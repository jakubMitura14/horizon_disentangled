from docx import Document

data = {
    "Full Name:": " Prof. Dr. Michael Kreißl",
    "Institution:": " University Hospital Magdeburg, Germany",
    "Position/Role in the project:": " Principal Investigator (PI)",
    "Email address:": " Not provided",
    "Phone number:": " Not provided",
}

doc = Document("grant_proposal/ProstateNET_Data_Access_Request_Form.docx")

instructions = {
    "1.1. Research Team": "Jakub Mitura, Joanna Wybrańska.",
    "2. Project Title/Study Name": "AI for Medical Imaging Optimization",
    "3. Research Question / Aim": "To develop AI tools and models that assist in the analysis and optimization of medical imaging workflows, ultimately enhancing diagnostic precision and supporting clinical decision-making.",
    "4. Scientific Background/Rationale": "Recent advancements in artificial intelligence offer significant opportunities to improve the interpretation and utility of complex medical datasets. This project aims to harness these capabilities to build robust models capable of analyzing multimodal clinical data. By focusing on fundamental representation learning and predictive modeling, we aim to provide actionable insights while ensuring the safety and transparency of the resulting systems.",
    "5. Methodology": "We will employ a multi-stage AI framework incorporating deep learning techniques, including representation learning and sequence modeling. The approach involves developing generalized supervisor models and exploring various neural architectures to effectively process and analyze the provided imaging datasets and corresponding clinical information.",
    "6. Dataset Requirements": "The project requires access to the entirety of the ProstateNET dataset, encompassing all available >17,000 cases and >1.5 million image representations, along with associated clinical and annotated data. Access to this complete, highly heterogeneous dataset across different vendors and clinical sites is vital. Training robust, vendor-agnostic foundation models and ensuring broad generalizability requires the maximum possible volume and diversity of data, making a partial subset insufficient for overcoming the current challenges in model generalizability.",
    "7. Expected Results and Applicability": "We expect to develop enhanced computational models capable of supporting diagnostic processes and offering quantitative insights into disease progression. The primary application is to serve as a decision-support mechanism for clinicians, providing probabilistic assessments to aid—but not replace—expert medical judgment.",
    "8. Funding Information": "Public Funding (University/Hospital)",
    "9. Study Duration": "36 Months (from May 1, 2025 to April 30, 2028).",
    "10. Ethical and Legal Considerations": "All research will be conducted using fully anonymized retrospective data, ensuring no personally identifiable information (PII) is accessed or processed. The study complies fully with GDPR requirements and relevant local regulations. Ethical approvals will be secured prior to any data processing.",
    "11. Supporting Documents": "PI CV uploaded."
}

def clean_insert(doc, title, new_text):
    for i, p in enumerate(doc.paragraphs):
        if p.text.startswith(title):
            # Target instruction
            target_idx = i
            while target_idx + 1 < len(doc.paragraphs):
                next_p = doc.paragraphs[target_idx+1]
                if next_p.text.strip() and not next_p.text[0].isdigit() and not next_p.style.name.startswith("Heading"):
                    target_idx += 1
                else:
                    break

            # Since the script modifies in place and can be run multiple times,
            # we need to cleanly check if it's already there
            search_idx = i + 1
            found = False
            while search_idx < len(doc.paragraphs):
                check_p = doc.paragraphs[search_idx]
                if check_p.text and check_p.text[0].isdigit() and (check_p.text[1] == '.' or (len(check_p.text) > 2 and check_p.text[2] == '.')):
                    break
                if check_p.text == new_text:
                    found = True
                    break
                search_idx += 1

            if not found:
                if target_idx + 1 < len(doc.paragraphs):
                    doc.paragraphs[target_idx + 1].insert_paragraph_before(new_text)
                else:
                    doc.add_paragraph(new_text)
            return

for p in doc.paragraphs:
    for key, val in data.items():
        if p.text.startswith(key):
            if not p.text.endswith(val):
                p.text = p.text.replace(key, key + val)

for title, text in instructions.items():
    clean_insert(doc, title, text)


# Clean up any leftover duplicated paragraphs from previous runs that were highly specific
text_to_remove = [
    "Leverage advanced generative architectures to synthesize high-fidelity medical imagery to enhance diagnostic precision while prioritizing patient safety, and modeling prostate cancer disease evolution with a causal AI framework.",
    "The project aims to develop a novel AI framework for managing prostate cancer by moving beyond simple prediction to a deep, causal understanding of disease progression. The goal is to create a transparent decision support tool that can integrate diverse, irregularly sampled clinical data (MRI, PET/CT, SPECT/CT, PSA, notes) and provide explainable insights through counterfactual reasoning. By disentangling the underlying biological signals from technical confounders, the model will simulate disease trajectories and treatment outcomes, mirroring a clinician's own reasoning process and fostering clinical trust.",
    "Multi-stage causal AI framework, generative models (Diffusion, VAEs, Transformers), Neural Jump ODEs for temporal trajectory modeling. The framework is built upon a sequential, multi-stage training process: Training Foundational Supervisor Models, Per-Modality Causal Representation Learning, Temporal Trajectory Modeling with Neural Jump ODEs, and Generative Synthesis and Clinical Outputs.",
    "~100,000 DICOM/NIfTI files (MRI, PET/CT, SPECT/CT). Dataset: ~300 TB (Raw/Processed). Stored in HDF5 format for efficient parallel I/O. Access to Exascale capabilities is vital for training high-fidelity generative models on this large dataset. Hybrid Julia/Python stack.",
    "High-fidelity synthetic image generation, counterfactual disease trajectories, radiation reduction. The AI system is designed as a decision support tool for medical professionals. Final diagnostic and treatment decisions remain with human doctors. The system provides probabilistic outputs and uncertainty estimates to aid, not replace, human judgment.",
    "Strictly synthetic and/or fully anonymized retrospective data, no PII, complies with GDPR and local regulations. Data usage is strictly for research purposes with appropriate ethics committee approval. Data is pseudonymized at source and anonymized for processing. Access is restricted to authorized personnel.",
    "Generative AI for Nuclear Medicine Optimization (CausalPCa)",
    "The project requires access to a substantial cohort of retrospective imaging studies (e.g., MRI, PET/CT) and associated clinical data. We anticipate utilizing approximately 100,000 multimodal imaging files, necessitating robust data handling and parallel processing capabilities to train advanced models effectively."
]

for p in list(doc.paragraphs):
    if p.text in text_to_remove:
        p.text = ""

# Keep only one instance of the names
seen_team = False
for p in list(doc.paragraphs):
    if p.text == "Jakub Mitura, Joanna Wybrańska.":
        if seen_team:
            p.text = ""
        else:
            seen_team = True

doc.save("grant_proposal/ProstateNET_Data_Access_Request_Form.docx")
