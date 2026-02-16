from docx import Document
import os

docx_path = "grant_proposal/latex/Regular Access - Project Scope and Plan.docx"

try:
    doc = Document(docx_path)
    text = []

    # Iterate over paragraphs
    for para in doc.paragraphs:
        if para.text.strip():
            text.append(f"PARA: {para.text.strip()}")

    # Iterate over tables
    for table in doc.tables:
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                if cell.text.strip():
                    row_text.append(cell.text.strip())
            if row_text:
                text.append(f"TABLE ROW: {' | '.join(row_text)}")

    with open("Regular_Access_Project_Scope_and_Plan.docx.txt", "w") as out:
        out.write("\n".join(text))
    print(f"Extracted content to Regular_Access_Project_Scope_and_Plan.docx.txt")
except Exception as e:
    print(f"Error extracting {docx_path}: {e}")
